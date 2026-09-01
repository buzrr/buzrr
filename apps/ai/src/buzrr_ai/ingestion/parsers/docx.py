"""DOCX parsing via python-docx.

Easier than PDF: Word carries real structure, so `Heading 1`..`Heading 9` style
names give the heading level directly with no heuristics.

DOCX has no page concept available without rendering, so every block reports
page 0 and citations fall back to the heading path alone.
"""

import re
from pathlib import Path

import docx

from buzrr_ai.ingestion.parsers.base import Block, ParsedDocument

_HEADING_STYLE = re.compile(r"^heading\s+(\d+)$", re.IGNORECASE)


def parse_docx(path: Path) -> ParsedDocument:
    document = docx.Document(str(path))
    blocks: list[Block] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        level = 0
        style_name = (paragraph.style.name if paragraph.style else "") or ""
        if match := _HEADING_STYLE.match(style_name.strip()):
            level = min(int(match.group(1)), 6)
        elif style_name.strip().lower() == "title":
            level = 1

        blocks.append(Block(text=text, page=0, heading_level=level))

    # Tables carry a lot of the substance in course material; flatten each row to
    # a pipe-delimited line so it survives into a chunk rather than vanishing.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(Block(text=" | ".join(cells), page=0, heading_level=0))

    return ParsedDocument(blocks=blocks, page_count=None)
